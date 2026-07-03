"""Document normalizer.

Поддерживаемые форматы:
- txt, md, csv (plain text)
- pdf (через pypdf)
- xlsx (через openpyxl, с table blocks)
- docx (через python-docx)

Для каждого файла создаётся NormalizedDocument с SourceSpan.
Таблицы (XLSX, CSV с заголовком) обрабатываются отдельно:
каждая ячейка = отдельный SourceSpan с row_index + column_name.
"""

from __future__ import annotations

import csv
import json
import logging
from pathlib import Path
from uuid import NAMESPACE_URL, uuid5

from scientific_tangle_gateway.schemas import (
    AccessLevel,
    DocumentArtifactPayload,
    DocumentBlockPayload,
    IngestionFilePayload,
    NormalizedDocumentPayload,
    SourceSpanPayload,
)

logger = logging.getLogger(__name__)


class DocumentNormalizer:
    def normalize_file(
        self,
        task_id: str,
        file_payload: IngestionFilePayload,
        access_level: AccessLevel,
    ) -> NormalizedDocumentPayload:
        source_path = Path(file_payload.storage_path)
        document_id = str(uuid5(NAMESPACE_URL, f"{task_id}:{source_path}"))
        source_type = self._source_type(source_path)
        parse_warnings: list[str] = []
        blocks: list[DocumentBlockPayload] = []
        tables: list[dict] = []
        source_spans: list[SourceSpanPayload] = []

        if file_payload.size_bytes == 0:
            parse_warnings.append("Файл пустой, текстовые блоки не созданы")
        elif source_type == "pdf":
            blocks.extend(self._read_pdf_blocks(document_id, source_path, parse_warnings))
        elif source_type == "xlsx":
            tables.extend(self._read_xlsx_tables(document_id, source_path, parse_warnings))
        elif source_type == "csv":
            tables.extend(self._read_csv_table(document_id, source_path, parse_warnings))
        elif source_type == "docx":
            blocks.extend(self._read_docx_blocks(document_id, source_path, parse_warnings))
        elif source_type in {"txt", "md"}:
            text = self._read_text_file(source_path, parse_warnings)
            if text:
                blocks.append(DocumentBlockPayload(
                    block_id=f"{document_id}:block:1",
                    text=text,
                ))
        else:
            parse_warnings.append(f"Формат {source_type} пока не поддержан normalizer")

        source_spans = self._build_source_spans(document_id, blocks, tables, access_level)
        if not source_spans and file_payload.size_bytes > 0:
            parse_warnings.append("Текст не извлечён, SourceSpan не создан")

        return NormalizedDocumentPayload(
            document_id=document_id,
            task_id=task_id,
            title=source_path.name,
            source_type=source_type,
            source_path=str(source_path),
            folder_category=source_path.parent.name,
            language=self._detect_language(blocks, tables),
            metadata={
                "content_type": file_payload.content_type,
                "size_bytes": file_payload.size_bytes,
                "original_filename": file_payload.filename,
            },
            access_level=access_level,
            blocks=blocks,
            tables=tables,
            source_spans=source_spans,
            parse_warnings=parse_warnings,
        )

    def _source_type(self, source_path: Path) -> str:
        suffix = source_path.suffix.lower().lstrip(".")
        return suffix or "unknown"

    def _read_text_file(self, source_path: Path, parse_warnings: list[str]) -> str:
        for encoding in ("utf-8-sig", "utf-8", "cp1251"):
            try:
                return source_path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        parse_warnings.append("Не удалось определить текстовую кодировку")
        return ""

    def _read_pdf_blocks(
        self,
        document_id: str,
        source_path: Path,
        parse_warnings: list[str],
    ) -> list[DocumentBlockPayload]:
        try:
            from pypdf import PdfReader
        except ImportError:
            parse_warnings.append("pypdf не установлен, PDF не обработан")
            return []

        try:
            reader = PdfReader(str(source_path))
        except Exception as error:
            parse_warnings.append(f"PDF не открыт: {error}")
            return []

        blocks = []
        for page_index, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception as error:
                parse_warnings.append(f"Страница {page_index}: текст не извлечён: {error}")
                continue

            text = text.strip()
            if text:
                blocks.append(
                    DocumentBlockPayload(
                        block_id=f"{document_id}:page:{page_index}",
                        page=page_index,
                        text=text,
                    )
                )

        if not blocks:
            parse_warnings.append("PDF не содержит извлекаемого текстового слоя")
        return blocks

    def _read_xlsx_tables(
        self,
        document_id: str,
        source_path: Path,
        parse_warnings: list[str],
    ) -> list[dict]:
        try:
            from openpyxl import load_workbook
        except ImportError:
            parse_warnings.append("openpyxl не установлен, XLSX не обработан")
            return []

        try:
            wb = load_workbook(filename=str(source_path), read_only=True, data_only=True)
        except Exception as error:
            parse_warnings.append(f"XLSX не открыт: {error}")
            return []

        tables = []
        for sheet_index, sheet_name in enumerate(wb.sheetnames, start=1):
            sheet = wb[sheet_name]
            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue
            header = [str(cell) if cell is not None else "" for cell in rows[0]]
            data_rows = []
            for row_idx, row in enumerate(rows[1:], start=2):
                if all(cell is None or cell == "" for cell in row):
                    continue
                data_rows.append({
                    "row_index": row_idx,
                    "cells": [str(cell) if cell is not None else "" for cell in row],
                })
            tables.append({
                "table_id": f"{document_id}:sheet:{sheet_index}",
                "sheet_name": sheet_name,
                "columns": header,
                "rows": data_rows,
                "row_count": len(data_rows),
            })
        wb.close()
        if not tables:
            parse_warnings.append("XLSX не содержит данных")
        return tables

    def _read_csv_table(
        self,
        document_id: str,
        source_path: Path,
        parse_warnings: list[str],
    ) -> list[dict]:
        text = self._read_text_file(source_path, parse_warnings)
        if not text:
            return []

        try:
            reader = csv.reader(text.splitlines())
            rows = list(reader)
        except Exception as error:
            parse_warnings.append(f"CSV не разобран: {error}")
            return []

        if len(rows) < 2:
            parse_warnings.append("CSV содержит менее 2 строк, как таблица не обработан")
            return []

        header = rows[0]
        data_rows = []
        for row_idx, row in enumerate(rows[1:], start=2):
            if not any(cell.strip() for cell in row):
                continue
            data_rows.append({
                "row_index": row_idx,
                "cells": [cell.strip() for cell in row],
            })
        return [{
            "table_id": f"{document_id}:csv:1",
            "sheet_name": source_path.stem,
            "columns": [c.strip() for c in header],
            "rows": data_rows,
            "row_count": len(data_rows),
        }]

    def _read_docx_blocks(
        self,
        document_id: str,
        source_path: Path,
        parse_warnings: list[str],
    ) -> list[DocumentBlockPayload]:
        try:
            from docx import Document
        except ImportError:
            parse_warnings.append("python-docx не установлен, DOCX не обработан")
            return []

        try:
            doc = Document(str(source_path))
        except Exception as error:
            parse_warnings.append(f"DOCX не открыт: {error}")
            return []

        blocks = []
        for para_idx, para in enumerate(doc.paragraphs, start=1):
            text = (para.text or "").strip()
            if text:
                blocks.append(
                    DocumentBlockPayload(
                        block_id=f"{document_id}:para:{para_idx}",
                        text=text,
                    )
                )

        for table_idx, table in enumerate(doc.tables, start=1):
            for row_idx, row in enumerate(table.rows, start=1):
                cells = [cell.text.strip() for cell in row.cells]
                blocks.append(
                    DocumentBlockPayload(
                        block_id=f"{document_id}:table{table_idx}:row{row_idx}",
                        text=" | ".join(cells),
                    )
                )

        if not blocks:
            parse_warnings.append("DOCX не содержит извлекаемого текста")
        return blocks

    def _build_source_spans(
        self,
        document_id: str,
        blocks: list[DocumentBlockPayload],
        tables: list[dict],
        access_level: AccessLevel,
    ) -> list[SourceSpanPayload]:
        spans: list[SourceSpanPayload] = []

        for index, block in enumerate(blocks, start=1):
            if not block.text:
                continue
            spans.append(
                SourceSpanPayload(
                    source_span_id=f"{document_id}:span:{index}",
                    document_id=document_id,
                    page=block.page,
                    char_start=0,
                    char_end=len(block.text),
                    raw_text=block.text,
                    parsed_text=block.text,
                    access_level=access_level,
                )
            )

        for table in tables:
            table_id = table["table_id"]
            columns = table.get("columns", [])
            for row in table.get("rows", []):
                row_index = row["row_index"]
                cells = row["cells"]
                for col_idx, cell_value in enumerate(cells):
                    if not cell_value:
                        continue
                    column_name = columns[col_idx] if col_idx < len(columns) else f"col_{col_idx + 1}"
                    spans.append(
                        SourceSpanPayload(
                            source_span_id=f"{table_id}:r{row_index}:c{col_idx + 1}",
                            document_id=document_id,
                            table_id=table_id,
                            row_index=row_index,
                            column_name=column_name,
                            char_start=0,
                            char_end=len(cell_value),
                            raw_text=cell_value,
                            parsed_text=cell_value,
                            access_level=access_level,
                        )
                    )

        return spans

    def _detect_language(
        self,
        blocks: list[DocumentBlockPayload],
        tables: list[dict],
    ) -> str:
        text_parts = [block.text for block in blocks]
        for table in tables:
            for row in table.get("rows", []):
                text_parts.extend(row.get("cells", []))
        text = " ".join(text_parts)
        if not text:
            return "unknown"
        cyrillic_count = sum(
            "а" <= char.lower() <= "я" or char.lower() == "ё" for char in text
        )
        return "ru" if cyrillic_count else "en"


class LocalNormalizedDocumentStore:
    def __init__(self, root_dir: Path) -> None:
        self.root_dir = root_dir

    def save_task_documents(
        self,
        task_id: str,
        documents: list[NormalizedDocumentPayload],
    ) -> list[DocumentArtifactPayload]:
        task_dir = self.root_dir / task_id
        task_dir.mkdir(parents=True, exist_ok=True)

        artifacts = []
        for document in documents:
            artifact_path = task_dir / f"{document.document_id}.json"
            artifact_path.write_text(
                json.dumps(
                    document.model_dump(mode="json"),
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )
            artifacts.append(self._to_artifact(document, artifact_path))
        return artifacts

    def list_task_documents(self, task_id: str) -> list[DocumentArtifactPayload]:
        task_dir = self.root_dir / task_id
        if not task_dir.exists():
            return []

        artifacts = []
        for artifact_path in sorted(task_dir.glob("*.json")):
            document = self._read_document(artifact_path)
            artifacts.append(self._to_artifact(document, artifact_path))
        return artifacts

    def get_document(self, document_id: str) -> NormalizedDocumentPayload | None:
        for artifact_path in self.root_dir.glob(f"*/{document_id}.json"):
            return self._read_document(artifact_path)
        return None

    def iter_all_documents(self) -> list[NormalizedDocumentPayload]:
        if not self.root_dir.exists():
            return []
        documents = []
        for artifact_path in sorted(self.root_dir.glob("*/*.json")):
            documents.append(self._read_document(artifact_path))
        return documents

    def _read_document(self, artifact_path: Path) -> NormalizedDocumentPayload:
        return NormalizedDocumentPayload.model_validate_json(
            artifact_path.read_text(encoding="utf-8")
        )

    def _to_artifact(
        self,
        document: NormalizedDocumentPayload,
        artifact_path: Path,
    ) -> DocumentArtifactPayload:
        return DocumentArtifactPayload(
            document_id=document.document_id,
            task_id=document.task_id,
            title=document.title,
            source_type=document.source_type,
            artifact_path=str(artifact_path),
            source_span_count=len(document.source_spans),
            parse_warnings=document.parse_warnings,
        )
